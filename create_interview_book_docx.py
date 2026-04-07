"""
Convert Interview Book from Markdown to DOCX

This script reads all markdown files from the interview book and creates
a formatted Microsoft Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path


def parse_markdown_to_docx(md_file: Path, doc: Document):
    """Parse a markdown file and add formatted content to the DOCX document."""

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_lang = ''
    code_lines = []
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start code block
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                # End code block
                in_code_block = False
                # Add code block to document
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Code'
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_data = []
            table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
            i += 1
            # Check if next line is separator
            if i < len(lines) and lines[i].strip().startswith('|') and '---' in lines[i]:
                i += 1  # Skip separator line
            continue
        else:
            if in_table:
                # End table, add to document
                if len(table_data) > 1:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            if row_idx == 0:  # Header row
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True

                    doc.add_paragraph()  # Space after table

                in_table = False
                table_data = []

        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            if level == 1:
                heading = doc.add_heading(text, level=1)
            elif level == 2:
                heading = doc.add_heading(text, level=2)
            elif level == 3:
                heading = doc.add_heading(text, level=3)
            else:
                heading = doc.add_heading(text, level=4)

            i += 1
            continue

        # Horizontal rules
        if line.strip() in ['---', '___', '***']:
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Blockquotes
        if line.startswith('>'):
            text = line[1:].strip()
            p = doc.add_paragraph(text)
            p.style = 'Intense Quote'
            i += 1
            continue

        # Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            # Parse inline formatting
            p = doc.add_paragraph()

            # Split by backticks for inline code
            parts = re.split(r'(`[^`]+`)', line)

            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    # Inline code
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    # Regular text with bold/italic
                    # Handle **bold**
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bold_part in bold_parts:
                        if bold_part.startswith('**') and bold_part.endswith('**'):
                            run = p.add_run(bold_part[2:-2])
                            run.bold = True
                        else:
                            # Handle *italic* or _italic_
                            italic_parts = re.split(r'(\*[^*]+\*|_[^_]+_)', bold_part)
                            for italic_part in italic_parts:
                                if (italic_part.startswith('*') and italic_part.endswith('*') and len(italic_part) > 2) or \
                                   (italic_part.startswith('_') and italic_part.endswith('_') and len(italic_part) > 2):
                                    run = p.add_run(italic_part[1:-1])
                                    run.italic = True
                                else:
                                    p.add_run(italic_part)

        i += 1


def create_interview_book_docx():
    """Create the complete interview book as a DOCX file."""

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Create Code style if it doesn't exist
    try:
        code_style = doc.styles['Code']
    except:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Title page
    title = doc.add_heading('Coding Interview Q&A Book', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('A Pattern-Focused Guide to Technical Interviews')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    tagline = doc.add_paragraph('Covering Arrays, Strings, Trees, Graphs, Dynamic Programming,')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline2 = doc.add_paragraph('Backtracking, and Binary Search Patterns')
    tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('Introduction', 'README.md'),
        ('Chapter 1: Core Data Structures (Easy)', None),
        ('  1.1 Arrays & Strings', 'chapter_1/01_arrays_strings.md'),
        ('  1.2 Hash Maps & Sets', 'chapter_1/02_hashmaps_sets.md'),
        ('  1.3 Trees & Graphs', 'chapter_1/03_trees_graphs.md'),
        ('  1.4 Recursion & Backtracking', 'chapter_1/04_recursion_backtracking.md'),
        ('  1.5 Sorting & Binary Search', 'chapter_1/05_sorting_binary_search.md'),
        ('  1.6 Dynamic Programming', 'chapter_1/06_dynamic_programming.md'),
        ('Chapter 2: Pattern Mastery (Medium)', None),
        ('  2.1 Sliding Window', 'chapter_2/01_sliding_window.md'),
        ('  2.2 Two Pointers', 'chapter_2/02_two_pointers.md'),
        ('  2.3 BFS & DFS', 'chapter_2/03_bfs_dfs_medium.md'),
        ('  2.4 Backtracking', 'chapter_2/04_backtracking_medium.md'),
        ('  2.5 Binary Search', 'chapter_2/05_binary_search_medium.md'),
        ('  2.6 Dynamic Programming', 'chapter_2/06_dynamic_programming_medium.md'),
        ('Chapter 3: Mock Interview Sessions', 'chapter_3/mock_sessions.md'),
    ]

    for item_text, _ in toc_items:
        doc.add_paragraph(item_text, style='List Bullet')

    doc.add_page_break()

    # Base path
    base_path = Path('/home/user/GenZ/interview_book')

    # Process each file
    files_to_process = [
        ('README.md', 'Introduction'),
        ('chapter_1/01_arrays_strings.md', 'Chapter 1.1: Arrays & Strings'),
        ('chapter_1/02_hashmaps_sets.md', 'Chapter 1.2: Hash Maps & Sets'),
        ('chapter_1/03_trees_graphs.md', 'Chapter 1.3: Trees & Graphs'),
        ('chapter_1/04_recursion_backtracking.md', 'Chapter 1.4: Recursion & Backtracking'),
        ('chapter_1/05_sorting_binary_search.md', 'Chapter 1.5: Sorting & Binary Search'),
        ('chapter_1/06_dynamic_programming.md', 'Chapter 1.6: Dynamic Programming'),
        ('chapter_2/01_sliding_window.md', 'Chapter 2.1: Sliding Window'),
        ('chapter_2/02_two_pointers.md', 'Chapter 2.2: Two Pointers'),
        ('chapter_2/03_bfs_dfs_medium.md', 'Chapter 2.3: BFS & DFS'),
        ('chapter_2/04_backtracking_medium.md', 'Chapter 2.4: Backtracking'),
        ('chapter_2/05_binary_search_medium.md', 'Chapter 2.5: Binary Search'),
        ('chapter_2/06_dynamic_programming_medium.md', 'Chapter 2.6: Dynamic Programming'),
        ('chapter_3/mock_sessions.md', 'Chapter 3: Mock Interview Sessions'),
    ]

    for file_path, chapter_title in files_to_process:
        print(f"Processing {file_path}...")

        # Add chapter heading
        doc.add_heading(chapter_title, level=0)

        # Parse markdown
        full_path = base_path / file_path
        if full_path.exists():
            parse_markdown_to_docx(full_path, doc)

        # Page break after each chapter
        doc.add_page_break()

    # Save document
    output_path = '/home/user/GenZ/Coding_Interview_QA_Book.docx'
    doc.save(output_path)
    print(f"\n✓ Book saved to: {output_path}")

    return output_path


if __name__ == "__main__":
    import sys

    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        print("✓ python-docx installed")

    # Create the book
    output = create_interview_book_docx()
    print(f"\n{'='*80}")
    print(f"INTERVIEW BOOK CREATED SUCCESSFULLY")
    print(f"{'='*80}")
    print(f"Output: {output}")
    print(f"Size: {Path(output).stat().st_size / 1024:.1f} KB")
    print(f"{'='*80}")
